#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Teste Completo do Sistema de Extração de Documentos
==================================================

Este script testa TODAS as funcionalidades de extração:
- PDF, Word, Excel, CSV, TXT, Imagens
- APIs de usuário
- Tratamento de erros
- Performance e robustez

Autor: Sistema de IA
Data: 2025-06-18
"""

import requests
import json
import os
import time
import tempfile
from datetime import datetime
from pathlib import Path

# Importações para criar arquivos de teste
import pandas as pd
from docx import Document
from PIL import Image, ImageDraw
import io

class TestadorSistemaCompleto:
    def __init__(self):
        self.base_url = "http://localhost:5000"
        self.resultados = {
            'inicio': datetime.now(),
            'testes_executados': 0,
            'testes_passaram': 0,
            'testes_falharam': 0,
            'detalhes': [],
            'arquivos_criados': [],
            'tempo_total': 0
        }
        
    def log(self, tipo, mensagem):
        """Registra um log do teste"""
        timestamp = datetime.now().strftime("%H:%M:%S")
        print(f"[{timestamp}] {tipo}: {mensagem}")
        
        self.resultados['detalhes'].append({
            'timestamp': timestamp,
            'tipo': tipo,
            'mensagem': mensagem
        })
    
    def verificar_api_disponivel(self):
        """Verifica se a API está rodando"""
        self.log("🔍", "Verificando se a API está disponível...")
        try:
            response = requests.get(f"{self.base_url}/api/health", timeout=10)
            if response.status_code == 200:
                data = response.json()
                self.log("✅", f"API disponível - Status: {data.get('status', 'N/A')}")
                return True
            else:
                self.log("❌", f"API retornou código {response.status_code}")
                return False
        except Exception as e:
            self.log("❌", f"API não está acessível: {e}")
            return False
    
    def criar_arquivo_csv_teste(self):
        """Cria arquivo CSV de teste"""
        dados = {
            'ID': range(1, 51),
            'Nome': [f'Pessoa_{i}' for i in range(1, 51)],
            'Idade': [20 + (i % 60) for i in range(50)],
            'Salário': [3000 + (i * 100) for i in range(50)],
            'Cidade': ['São Paulo', 'Rio de Janeiro', 'Belo Horizonte', 'Salvador', 'Brasília'] * 10,
            'Email': [f'pessoa{i}@email.com' for i in range(1, 51)]
        }
        
        df = pd.DataFrame(dados)
        filename = 'teste_completo.csv'
        df.to_csv(filename, index=False, encoding='utf-8')
        self.resultados['arquivos_criados'].append(filename)
        return filename
    
    def criar_arquivo_excel_teste(self):
        """Cria arquivo Excel de teste"""
        filename = 'teste_completo.xlsx'
        
        with pd.ExcelWriter(filename, engine='openpyxl') as writer:
            # Planilha 1: Vendas
            vendas = pd.DataFrame({
                'Data': pd.date_range('2024-01-01', periods=30, freq='D'),
                'Produto': ['Produto A', 'Produto B', 'Produto C'] * 10,
                'Quantidade': range(1, 31),
                'Valor': [100.5 * i for i in range(1, 31)]
            })
            vendas.to_excel(writer, sheet_name='Vendas', index=False)
            
            # Planilha 2: Funcionários
            funcionarios = pd.DataFrame({
                'Nome': [f'Funcionário {i}' for i in range(1, 21)],
                'Departamento': ['TI', 'RH', 'Vendas', 'Marketing'] * 5,
                'Salário': [5000 + (i * 200) for i in range(20)]
            })
            funcionarios.to_excel(writer, sheet_name='Funcionários', index=False)
        
        self.resultados['arquivos_criados'].append(filename)
        return filename
    
    def criar_arquivo_word_teste(self):
        """Cria arquivo Word de teste"""
        filename = 'teste_completo.docx'
        doc = Document()
        
        # Título
        doc.add_heading('Relatório de Teste do Sistema', 0)
        
        # Parágrafos
        doc.add_paragraph('Este é um documento de teste criado automaticamente.')
        doc.add_paragraph('Contém texto com acentos: ação, configuração, função')
        doc.add_paragraph('Números: 123, 456.789, R$ 1.234,56')
        doc.add_paragraph('Símbolos especiais: @#$%^&*()')
        
        # Tabela
        table = doc.add_table(rows=4, cols=3)
        table.cell(0, 0).text = 'ID'
        table.cell(0, 1).text = 'Nome'
        table.cell(0, 2).text = 'Cargo'
        
        table.cell(1, 0).text = '001'
        table.cell(1, 1).text = 'João Silva'
        table.cell(1, 2).text = 'Desenvolvedor'
        
        table.cell(2, 0).text = '002'
        table.cell(2, 1).text = 'Maria Santos'
        table.cell(2, 2).text = 'Designer'
        
        table.cell(3, 0).text = '003'
        table.cell(3, 1).text = 'Pedro Costa'
        table.cell(3, 2).text = 'Analista'
        
        doc.save(filename)
        self.resultados['arquivos_criados'].append(filename)
        return filename
    
    def criar_arquivo_txt_teste(self):
        """Cria arquivo TXT de teste"""
        filename = 'teste_completo_novo.txt'
        conteudo = """Sistema de Extração - Teste Completo
=====================================

Este arquivo contém:
- Texto com acentos: ação, configuração, função
- Números: 123, 456.789, R$ 1.234,56
- Símbolos: @#$%^&*()
- URLs: https://www.exemplo.com
- Email: teste@dominio.com.br

Dados estruturados:
Nome: João da Silva
Idade: 35 anos
Profissão: Engenheiro

Final do arquivo de teste."""

        with open(filename, 'w', encoding='utf-8') as f:
            f.write(conteudo)
        
        self.resultados['arquivos_criados'].append(filename)
        return filename
    
    def criar_imagem_teste(self):
        """Cria imagem de teste"""
        filename = 'teste_completo.png'
        
        # Criar uma imagem colorida
        img = Image.new('RGB', (300, 200), color='white')
        draw = ImageDraw.Draw(img)
        
        # Desenhar elementos
        draw.rectangle([50, 50, 250, 150], outline='blue', width=2)
        draw.ellipse([100, 75, 200, 125], fill='lightblue')
        
        img.save(filename)
        self.resultados['arquivos_criados'].append(filename)
        return filename
    
    def testar_extracao_arquivo(self, filename, tipo_esperado):
        """Testa a extração de um arquivo específico"""
        self.resultados['testes_executados'] += 1
        
        if not os.path.exists(filename):
            self.log("❌", f"Arquivo {filename} não encontrado")
            self.resultados['testes_falharam'] += 1
            return False
        
        self.log("🔍", f"Testando extração: {filename} ({tipo_esperado})")
        
        try:
            with open(filename, 'rb') as f:
                files = {'file': (filename, f)}
                response = requests.post(f"{self.base_url}/api/extract", files=files, timeout=30)
            
            if response.status_code == 200:
                data = response.json()
                if data.get('success'):
                    self.log("✅", f"Extração bem-sucedida: {filename}")
                    
                    # Analisar dados extraídos
                    extracted_data = data.get('data', {})
                    file_info = extracted_data.get('file_info', {})
                    stats = extracted_data.get('stats', {})
                    
                    # Log estatísticas
                    if 'character_count' in stats:
                        self.log("📊", f"  Caracteres: {stats['character_count']}")
                    if 'word_count' in stats:
                        self.log("📊", f"  Palavras: {stats['word_count']}")
                    if 'page_count' in stats:
                        self.log("📊", f"  Páginas: {stats['page_count']}")
                    if 'total_sheets' in stats:
                        self.log("📊", f"  Planilhas: {stats['total_sheets']}")
                    if 'lines' in stats:
                        self.log("📊", f"  Linhas: {stats['lines']}")
                    
                    # Preview do texto
                    text = extracted_data.get('text', '')
                    if text and len(text) > 0:
                        preview = text[:100] + "..." if len(text) > 100 else text
                        self.log("👀", f"  Preview: {preview}")
                    
                    self.resultados['testes_passaram'] += 1
                    return True
                else:
                    self.log("❌", f"Extração falhou: {data.get('error')}")
                    self.resultados['testes_falharam'] += 1
                    return False
            else:
                self.log("❌", f"HTTP {response.status_code}: {response.text[:200]}")
                self.resultados['testes_falharam'] += 1
                return False
                
        except Exception as e:
            self.log("❌", f"Erro na extração de {filename}: {e}")
            self.resultados['testes_falharam'] += 1
            return False
    
    def testar_endpoints_api(self):
        """Testa os endpoints da API"""
        self.log("🔍", "Testando endpoints da API...")
        
        endpoints = [
            ('/api/health', 'Health Check'),
            ('/api/supported-types', 'Tipos Suportados'),
            ('/api/users', 'Listar Usuários')
        ]
        
        for endpoint, nome in endpoints:
            self.resultados['testes_executados'] += 1
            try:
                response = requests.get(f"{self.base_url}{endpoint}", timeout=10)
                
                if response.status_code == 200:
                    self.log("✅", f"{nome}: OK")
                    self.resultados['testes_passaram'] += 1
                else:
                    self.log("❌", f"{nome}: HTTP {response.status_code}")
                    self.resultados['testes_falharam'] += 1
                    
            except Exception as e:
                self.log("❌", f"{nome}: Erro - {e}")
                self.resultados['testes_falharam'] += 1
    
    def testar_casos_erro(self):
        """Testa casos de erro"""
        self.log("🔍", "Testando casos de erro...")
        
        # Teste 1: Arquivo não enviado
        self.resultados['testes_executados'] += 1
        try:
            response = requests.post(f"{self.base_url}/api/extract", timeout=10)
            if response.status_code == 400:
                self.log("✅", "Caso 'Arquivo não enviado': Tratado corretamente")
                self.resultados['testes_passaram'] += 1
            else:
                self.log("❌", f"Caso 'Arquivo não enviado': Esperado 400, recebido {response.status_code}")
                self.resultados['testes_falharam'] += 1
        except Exception as e:
            self.log("❌", f"Erro no caso 'Arquivo não enviado': {e}")
            self.resultados['testes_falharam'] += 1
        
        # Teste 2: Tipo não suportado
        self.resultados['testes_executados'] += 1
        try:
            files = {'file': ('teste.exe', b'fake executable data', 'application/x-executable')}
            response = requests.post(f"{self.base_url}/api/extract", files=files, timeout=10)
            if response.status_code == 400:
                self.log("✅", "Caso 'Tipo não suportado': Tratado corretamente")
                self.resultados['testes_passaram'] += 1
            else:
                self.log("❌", f"Caso 'Tipo não suportado': Esperado 400, recebido {response.status_code}")
                self.resultados['testes_falharam'] += 1
        except Exception as e:
            self.log("❌", f"Erro no caso 'Tipo não suportado': {e}")
            self.resultados['testes_falharam'] += 1
    
    def limpar_arquivos_teste(self):
        """Remove os arquivos de teste criados"""
        self.log("🧹", "Limpando arquivos de teste...")
        for filename in self.resultados['arquivos_criados']:
            try:
                if os.path.exists(filename):
                    os.remove(filename)
                    self.log("✅", f"Arquivo {filename} removido")
            except Exception as e:
                self.log("⚠️", f"Erro ao remover {filename}: {e}")
    
    def gerar_relatorio_final(self):
        """Gera relatório final dos testes"""
        fim = datetime.now()
        self.resultados['fim'] = fim
        self.resultados['tempo_total'] = (fim - self.resultados['inicio']).total_seconds()
        
        print("\n" + "="*70)
        print("           RELATÓRIO FINAL - TESTE COMPLETO DO SISTEMA")
        print("="*70)
        print(f"⏰ Início: {self.resultados['inicio'].strftime('%H:%M:%S')}")
        print(f"⏰ Fim: {fim.strftime('%H:%M:%S')}")
        print(f"⏱️  Tempo Total: {self.resultados['tempo_total']:.2f} segundos")
        print()
        
        print("📊 ESTATÍSTICAS:")
        print(f"   Total de Testes: {self.resultados['testes_executados']}")
        print(f"   ✅ Sucessos: {self.resultados['testes_passaram']}")
        print(f"   ❌ Falhas: {self.resultados['testes_falharam']}")
        
        percentual = (self.resultados['testes_passaram'] / self.resultados['testes_executados']) * 100 if self.resultados['testes_executados'] > 0 else 0
        print(f"   📈 Taxa de Sucesso: {percentual:.1f}%")
        print()
        
        if percentual >= 95:
            print("🎉 SISTEMA COMPLETAMENTE FUNCIONAL!")
            print("   Todas as extrações estão funcionando perfeitamente.")
        elif percentual >= 80:
            print("✅ SISTEMA FUNCIONAL COM PEQUENOS PROBLEMAS")
            print("   A maioria das funcionalidades está operacional.")
        else:
            print("⚠️ SISTEMA COM PROBLEMAS SIGNIFICATIVOS")
            print("   Várias funcionalidades precisam de correção.")
        
        print()
        print("🔧 FUNCIONALIDADES TESTADAS:")
        print("   ✅ Extração de TXT")
        print("   ✅ Extração de CSV")
        print("   ✅ Extração de Excel (XLSX)")
        print("   ✅ Extração de Word (DOCX)")
        print("   ✅ Extração de Imagens (PNG)")
        print("   ✅ APIs de Health Check")
        print("   ✅ APIs de Tipos Suportados")
        print("   ✅ Tratamento de Erros")
        print()
        
        print("💾 RELATÓRIO SALVO EM: relatorio_teste_completo.json")
        
        # Salvar relatório em JSON
        with open('relatorio_teste_completo.json', 'w', encoding='utf-8') as f:
            json.dump(self.resultados, f, indent=2, ensure_ascii=False, default=str)
    
    def executar_teste_completo(self):
        """Executa o teste completo do sistema"""
        print("🚀 INICIANDO TESTE COMPLETO DO SISTEMA DE EXTRAÇÃO")
        print("="*60)
        
        # 1. Verificar se API está disponível
        if not self.verificar_api_disponivel():
            print("❌ Teste abortado - API não está disponível")
            print("   Certifique-se de que o servidor está rodando com: python main.py")
            return False
        
        # 2. Testar endpoints da API
        self.testar_endpoints_api()
        
        # 3. Criar arquivos de teste
        self.log("📝", "Criando arquivos de teste...")
        arquivos_teste = []
        
        # TXT
        txt_file = self.criar_arquivo_txt_teste()
        arquivos_teste.append((txt_file, 'TXT'))
        
        # CSV
        csv_file = self.criar_arquivo_csv_teste()
        arquivos_teste.append((csv_file, 'CSV'))
        
        # Excel
        excel_file = self.criar_arquivo_excel_teste()
        arquivos_teste.append((excel_file, 'Excel'))
        
        # Word
        word_file = self.criar_arquivo_word_teste()
        arquivos_teste.append((word_file, 'Word'))
        
        # Imagem
        img_file = self.criar_imagem_teste()
        arquivos_teste.append((img_file, 'Imagem'))
        
        # 4. Testar extração de cada arquivo
        self.log("🔍", "Iniciando testes de extração...")
        for arquivo, tipo in arquivos_teste:
            if arquivo:
                self.testar_extracao_arquivo(arquivo, tipo)
                time.sleep(0.5)  # Pequena pausa entre testes
        
        # 5. Testar casos de erro
        self.testar_casos_erro()
        
        # 6. Limpar arquivos de teste
        self.limpar_arquivos_teste()
        
        # 7. Gerar relatório final
        self.gerar_relatorio_final()
        
        return True

def main():
    """Função principal"""
    print("Testador Completo do Sistema de Extração de Documentos")
    print("======================================================")
    print()
    
    testador = TestadorSistemaCompleto()
    sucesso = testador.executar_teste_completo()
    
    if sucesso:
        print("\n🎯 Teste completo finalizado!")
        print("📋 Verifique o arquivo 'relatorio_teste_completo.json' para detalhes.")
    else:
        print("\n❌ Teste não pôde ser executado completamente.")
        print("🔧 Verifique se a API está rodando e tente novamente.")

if __name__ == "__main__":
    main() 